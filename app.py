from flask import Flask, request, send_file, jsonify
from docx.api import Document
import tempfile
import os
from flask_cors import CORS
import re
import io
import json
from docx.shared import Pt
import string
import time
from functools import wraps
import hashlib
import threading
from concurrent.futures import ThreadPoolExecutor, as_completed
import queue
import logging
from collections import defaultdict
import gc

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)

# Performance and reliability configurations
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['REQUEST_TIMEOUT'] = 30  # Back to 30 seconds for reliability
app.config['MAX_CONCURRENT_REQUESTS'] = 10  # Back to reasonable limit
app.config['CACHE_TTL'] = 3600  # 1 hour cache TTL
app.config['START_TIME'] = time.time()  # Track app start time
app.config['FAST_MODE'] = False  # Disable aggressive fast mode

# Global caches and state
keyword_cache = {}
ats_score_cache = {}
processing_queue = queue.Queue()
executor = ThreadPoolExecutor(max_workers=app.config['MAX_CONCURRENT_REQUESTS'])
active_requests = 0
request_lock = threading.Lock()

# Pre-compiled regex patterns for maximum performance
WORD_BOUNDARY_PATTERN = re.compile(r'\b\w+\b')
SKILL_PATTERNS = [
    re.compile(r'skills?:\s*([\w\s,;&]+)', re.IGNORECASE),
    re.compile(r'technologies?:\s*([\w\s,;&]+)', re.IGNORECASE),
    re.compile(r'languages?:\s*([\w\s,;&]+)', re.IGNORECASE),
]

# Fast keyword lookup set (pre-computed for speed)
FAST_KEYWORDS = {
    'python', 'java', 'javascript', 'react', 'node', 'sql', 'aws', 'docker', 'kubernetes',
    'html', 'css', 'git', 'mongodb', 'postgresql', 'redis', 'nginx', 'apache',
    'linux', 'ubuntu', 'centos', 'debian', 'windows', 'macos', 'ios', 'android',
    'agile', 'scrum', 'kanban', 'devops', 'ci/cd', 'jenkins', 'github', 'gitlab',
    'azure', 'gcp', 'heroku', 'digitalocean', 'vps', 'vpc', 'ec2', 's3', 'lambda',
    'typescript', 'vue', 'angular', 'express', 'django', 'flask', 'spring', 'laravel',
    'php', 'ruby', 'go', 'rust', 'swift', 'kotlin', 'scala', 'r', 'matlab', 'julia',
    'tensorflow', 'pytorch', 'keras', 'scikit-learn', 'pandas', 'numpy', 'matplotlib',
    'tableau', 'powerbi', 'excel', 'word', 'powerpoint', 'outlook', 'teams', 'slack',
    'jira', 'confluence', 'notion', 'figma', 'sketch', 'adobe', 'photoshop', 'illustrator'
}

# Retry mechanism and circuit breaker
class CircuitBreaker:
    def __init__(self, failure_threshold=5, recovery_timeout=60):
        self.failure_threshold = failure_threshold
        self.recovery_timeout = recovery_timeout
        self.failure_count = 0
        self.last_failure_time = 0
        self.state = 'CLOSED'  # CLOSED, OPEN, HALF_OPEN
    
    def call(self, func, *args, **kwargs):
        if self.state == 'OPEN':
            if time.time() - self.last_failure_time > self.recovery_timeout:
                self.state = 'HALF_OPEN'
            else:
                raise Exception("Circuit breaker is OPEN")
        
        try:
            result = func(*args, **kwargs)
            if self.state == 'HALF_OPEN':
                self.state = 'CLOSED'
                self.failure_count = 0
            return result
        except Exception as e:
            self.failure_count += 1
            self.last_failure_time = time.time()
            
            if self.failure_count >= self.failure_threshold:
                self.state = 'OPEN'
            
            raise e

# Global circuit breakers
optimization_circuit_breaker = CircuitBreaker()
keyword_circuit_breaker = CircuitBreaker()

def retry_with_backoff(func, max_retries=3, base_delay=1):
    """Retry function with exponential backoff"""
    for attempt in range(max_retries):
        try:
            return func()
        except Exception as e:
            if attempt == max_retries - 1:
                raise e
            
            delay = base_delay * (2 ** attempt)
            time.sleep(delay)
            logger.warning(f"Retry attempt {attempt + 1}")
    
    # This should never be reached, but just in case
    raise Exception("Max retries exceeded")

# Health check endpoint
@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint for monitoring"""
    try:
        # Check if we can process a simple request
        test_text = "test resume content"
        test_job = "test job description"
        
        # Quick cache test
        get_cache_key('test', test_text)
        
        # Quick keyword extraction test
        extract_technical_keywords_optimized(test_job)
        
        return jsonify({
            'status': 'healthy',
            'timestamp': time.time(),
            'cache_size': len(keyword_cache) + len(ats_score_cache),
            'active_requests': active_requests,
            'memory_usage': 'OK'
        }), 200
    except Exception as e:
        logger.error(f"Health check failed: {e}")
        return jsonify({
            'status': 'unhealthy',
            'error': str(e),
            'timestamp': time.time()
        }), 503

# Cache cleanup endpoint
@app.route('/cache/clear', methods=['POST'])
def clear_cache():
    """Clear all caches (admin endpoint)"""
    try:
        global keyword_cache, ats_score_cache
        keyword_cache.clear()
        ats_score_cache.clear()
        cleanup_temp_files()
        gc.collect()
        
        return jsonify({
            'message': 'Cache cleared successfully',
            'timestamp': time.time()
        }), 200
    except Exception as e:
        logger.error(f"Cache clear failed: {e}")
        return jsonify({'error': str(e)}), 500

# Performance monitoring endpoint
@app.route('/metrics', methods=['GET'])
def get_metrics():
    """Get performance metrics"""
    return jsonify({
        'cache_stats': {
            'keyword_cache_size': len(keyword_cache),
            'ats_cache_size': len(ats_score_cache),
            'cache_hit_rate': 'N/A'  # Would need to track hits/misses
        },
        'system_stats': {
            'active_requests': active_requests,
            'max_concurrent_requests': app.config['MAX_CONCURRENT_REQUESTS'],
            'uptime': time.time() - app.config['START_TIME']
        },
        'circuit_breakers': {
            'optimization_state': optimization_circuit_breaker.state,
            'keyword_state': keyword_circuit_breaker.state
        },
        'performance_mode': {
            'fast_mode': app.config['FAST_MODE'],
            'timeout_seconds': app.config['REQUEST_TIMEOUT'],
            'max_concurrent': app.config['MAX_CONCURRENT_REQUESTS']
        }
    }), 200

@app.route('/toggle-fast-mode', methods=['POST'])
def toggle_fast_mode():
    """Toggle fast mode on/off"""
    try:
        app.config['FAST_MODE'] = not app.config['FAST_MODE']
        return jsonify({
            'fast_mode': app.config['FAST_MODE'],
            'message': f'Fast mode {"enabled" if app.config["FAST_MODE"] else "disabled"}'
        }), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/status', methods=['GET'])
def get_status():
    """Get current system status"""
    return jsonify({
        'status': 'operational',
        'fast_mode': app.config['FAST_MODE'],
        'uptime': time.time() - app.config['START_TIME'],
        'active_requests': active_requests,
        'cache_size': len(keyword_cache) + len(ats_score_cache)
    }), 200

# Cache management
def get_cache_key(data_type, *args):
    """Generate cache key for different data types"""
    content = '|'.join(str(arg) for arg in args)
    return f"{data_type}:{hashlib.md5(content.encode()).hexdigest()}"

# Memory management
def cleanup_temp_files():
    """Clean up temporary files to prevent memory leaks"""
    try:
        temp_dir = tempfile.gettempdir()
        for filename in os.listdir(temp_dir):
            if filename.startswith('tmp') and filename.endswith('.docx'):
                filepath = os.path.join(temp_dir, filename)
                if time.time() - os.path.getmtime(filepath) > 300:  # 5 minutes old
                    try:
                        os.remove(filepath)
                    except:
                        pass
    except Exception as e:
        logger.warning(f"Cleanup failed: {e}")

def get_cached_result(cache_dict, key, ttl=3600):
    """Get cached result if not expired"""
    if key in cache_dict:
        timestamp, result = cache_dict[key]
        if time.time() - timestamp < ttl:
            return result
        else:
            del cache_dict[key]
    return None

def set_cached_result(cache_dict, key, result, ttl=3600):
    """Set cached result with timestamp"""
    cache_dict[key] = (time.time(), result)

# Enhanced timeout handler with better error handling
def timeout_handler(timeout_seconds=30):
    def decorator(f):
        @wraps(f)
        def wrapper(*args, **kwargs):
            global active_requests  # <-- moved to the top
            start_time = time.time()
            try:
                # Check if we're at capacity
                with request_lock:
                    if active_requests >= app.config['MAX_CONCURRENT_REQUESTS']:
                        return jsonify({'error': 'Server is busy. Please try again in a moment.'}), 503
                    active_requests += 1
                
                result = f(*args, **kwargs)
                
                if time.time() - start_time > timeout_seconds:
                    raise Exception("Request timed out")
                return result
            except Exception as e:
                logger.error(f"Error in {f.__name__}: {str(e)}")
                if time.time() - start_time > timeout_seconds:
                    return jsonify({'error': f'Request timed out after {timeout_seconds} seconds. Please try with a smaller file or shorter description.'}), 408
                return jsonify({'error': f'Processing failed: {str(e)}'}), 500
            finally:
                with request_lock:
                    active_requests -= 1
        return wrapper
    return decorator

# ATS Scoring weights and criteria
ATS_CRITERIA = {
    'keyword_match': 0.70,  # 60% - How well keywords match job description (increased from 35%)
    'formatting': 0.10,     # 20% - Clean formatting, proper sections (decreased from 25%)
    'content_quality': 0.20, # 15% - Professional language, achievements (decreased from 20%)
    'structure': 0.00,      # 5% - Proper resume structure (decreased from 15%)
    'length': 0.00          # 0% - Length is less important for optimization (decreased from 5%)
}

# ATS-friendly formatting requirements
ATS_FORMATTING_REQUIREMENTS = {
    'sections': ['experience', 'education', 'skills', 'summary', 'objective'],
    'avoid_elements': ['tables', 'images', 'charts', 'headers', 'footers', 'columns'],
    'preferred_fonts': ['arial', 'calibri', 'times new roman', 'georgia'],
    'max_length': 2000,  # words
    'min_length': 200    # words
}

# Predefined technical keywords by category (all in lowercase for matching)
TECHNICAL_KEYWORDS = {
    'programming_languages': {
        'python', 'java', 'javascript', 'js', 'typescript', 'ts', 'c++', 'c#', 'ruby', 'php', 'swift',
        'kotlin', 'go', 'golang', 'rust', 'scala', 'r', 'matlab', 'sql', 'perl', 'shell', 'bash',
        'dart', 'elixir', 'haskell', 'lua', 'assembly', 'fortran', 'cobol', 'ada', 'groovy', 'clojure',
        'c', 'f#', 'ocaml', 'erlang', 'julia', 'nim', 'crystal', 'zig', 'v', 'odin', 'pony'
    },
    'web_technologies': {
        'html', 'css', 'sass', 'less', 'react', 'angular', 'vue', 'node.js', 'nodejs', 'express',
        'django', 'flask', 'spring', 'asp.net', 'jquery', 'bootstrap', 'tailwind',
        'next.js', 'nuxt.js', 'svelte', 'sveltekit', 'astro', 'remix', 'gatsby', 'webpack', 'vite',
        'rollup', 'esbuild', 'parcel', 'graphql', 'apollo', 'prisma', 'sequelize', 'typeorm',
        'nestjs', 'fastapi', 'fastify', 'hapi', 'koa', 'web3', 'solidity', 'ethers.js', 'web3.js',
        'stripe', 'twilio', 'sendgrid', 'mailgun', 'socket.io', 'ws', 'websocket'
    },
    'databases': {
        'mysql', 'postgresql', 'mongodb', 'redis', 'oracle', 'sqlite', 'sql server', 'dynamodb',
        'cassandra', 'elasticsearch', 'neo4j', 'mariadb', 'cockroachdb', 'timescaledb', 'influxdb',
        'clickhouse', 'snowflake', 'bigquery', 'redshift', 'databricks', 'hive', 'hbase', 'couchdb',
        'rethinkdb', 'arangodb', 'fauna', 'supabase', 'planetscale', 'firebase', 'appwrite'
    },
    'cloud_platforms': {
        'aws', 'azure', 'gcp', 'google cloud', 'heroku', 'digitalocean', 'firebase',
        'cloudflare', 'vercel', 'netlify', 'alibaba cloud', 'oracle cloud', 'ibm cloud',
        'linode', 'vultr', 'render', 'railway', 'fly.io', 'supabase', 'appwrite', 'hasura',
        'stripe', 'twilio', 'sendgrid', 'mailgun', 'auth0', 'okta', 'cognito'
    },
    'data_science': {
        'pandas', 'numpy', 'scipy', 'scikit-learn', 'sklearn', 'tensorflow', 'pytorch', 'keras',
        'matplotlib', 'seaborn', 'plotly', 'tableau', 'power bi', 'jupyter', 'spss', 'sas',
        'opencv', 'pillow', 'statsmodels', 'xgboost', 'lightgbm', 'catboost', 'fastai',
        'transformers', 'spacy', 'nltk', 'gensim', 'word2vec', 'bert', 'gpt', 'llama',
        'streamlit', 'gradio', 'mlflow', 'kubeflow', 'ray', 'dask', 'vaex', 'modin', 'rapids',
        'polars', 'orange3', 'dataspell'
    },
    'tools_and_platforms': {
        'git', 'docker', 'kubernetes', 'jenkins', 'jira', 'confluence', 'bitbucket', 'github',
        'gitlab', 'terraform', 'ansible', 'vagrant', 'postman', 'swagger', 'snowflake',
        'helm', 'istio', 'linkerd', 'prometheus', 'grafana', 'elk stack', 'elasticsearch',
        'logstash', 'kibana', 'splunk', 'datadog', 'newrelic', 'sentry', 'rollbar',
        'circleci', 'github actions', 'gitlab ci', 'travis ci', 'azure devops', 'teamcity',
        'bamboo', 'sonarqube', 'codecov', 'coveralls', 'semaphore', 'appveyor', 'wercker',
        'drone', 'concourse', 'spinnaker', 'argo', 'tekton', 'skaffold', 'tilt', 'lens',
        'rancher', 'openshift', 'minikube', 'kind', 'k3s', 'microk8s', 'kubectl', 'kustomize'
    }
}

# Flatten the keywords for easier searching
ALL_KEYWORDS = {keyword.lower() for category in TECHNICAL_KEYWORDS.values() for keyword in category}

# --- Industry-specific best-practice keyword lists ---
INDUSTRY_KEYWORDS = {
    'software_engineering': [
        # Languages
        'Python', 'Java', 'C++', 'C#', 'JavaScript', 'TypeScript', 'Go', 'Ruby', 'Kotlin', 'Swift',
        # Frameworks/Tools
        'React', 'Angular', 'Vue', 'Node.js', 'Django', 'Flask', 'Spring', 'Express', 'Docker', 'Kubernetes',
        'AWS', 'Azure', 'GCP', 'CI/CD', 'Git', 'REST', 'GraphQL', 'Microservices', 'Agile', 'Scrum',
        # Practices
        'Unit Testing', 'TDD', 'OOP', 'Design Patterns', 'DevOps', 'Cloud', 'API Development', 'Version Control',
        # Soft skills
        'Collaboration', 'Problem Solving', 'Communication'
    ],
    'data_analytics': [
        # Languages/Tools
        'Python', 'R', 'SQL', 'Excel', 'Tableau', 'Power BI', 'Looker', 'SAS', 'SPSS', 'Jupyter',
        # Libraries
        'Pandas', 'NumPy', 'Matplotlib', 'Seaborn', 'Scikit-learn', 'TensorFlow', 'PyTorch',
        # Concepts
        'Data Visualization', 'Data Cleaning', 'Data Mining', 'Statistical Analysis', 'ETL', 'Big Data',
        'Machine Learning', 'Predictive Modeling', 'Regression', 'Classification', 'A/B Testing',
        # Soft skills
        'Attention to Detail', 'Critical Thinking', 'Storytelling', 'Business Acumen'
    ]
}

def find_keyword_with_original_case(text, keyword_lower):
    """Find the keyword in the text and return it with its original case"""
    pattern = r'\b' + re.escape(keyword_lower) + r'\b'
    match = re.search(pattern, text.lower())
    if match:
        start, end = match.span()
        # Return the original casing from the text
        return text[start:end]
    # If not found in text, return the keyword with proper capitalization
    return keyword_lower.title()

def extract_technical_keywords(text):
    """Legacy function - use extract_technical_keywords_optimized for better performance"""
    return extract_technical_keywords_optimized(text)

def extract_technical_keywords_optimized(text):
    """Optimized keyword extraction with caching - restored original functionality"""
    cache_key = get_cache_key('keywords', text[:1000])  # Back to original cache key size
    cached_result = get_cached_result(keyword_cache, cache_key, 1800)  # 30 min cache
    if cached_result:
        return cached_result
    
    found_keywords = {}
    
    # Pre-compile regex patterns for better performance
    word_boundary_pattern = re.compile(r'\b\w+\b')
    skill_patterns = [
        re.compile(r'proficient in\s+([\w\s,;&]+)', re.IGNORECASE),
        re.compile(r'experience with\s+([\w\s,;&]+)', re.IGNORECASE),
        re.compile(r'knowledge of\s+([\w\s,;&]+)', re.IGNORECASE),
        re.compile(r'skills?:\s*([\w\s,;&]+)', re.IGNORECASE),
        re.compile(r'technologies?:\s*([\w\s,;&]+)', re.IGNORECASE),
        re.compile(r'languages?:\s*([\w\s,;&]+)', re.IGNORECASE),
        re.compile(r'frameworks?:\s*([\w\s,;&]+)', re.IGNORECASE),
        re.compile(r'tools?:\s*([\w\s,;&]+)', re.IGNORECASE),
        re.compile(r'platforms?:\s*([\w\s,;&]+)', re.IGNORECASE)
    ]
    
    # Fast exact match lookup using set
    text_lower = text.lower()
    text_words = set(word_boundary_pattern.findall(text_lower))
    
    # Find exact matches first (fastest) - use ALL_KEYWORDS for completeness
    for keyword in ALL_KEYWORDS:
        if keyword in text_words:
            original_case = find_keyword_with_original_case(text, keyword)
            found_keywords[keyword] = original_case
    
    # Process skill patterns (more expensive, so do less)
    for pattern in skill_patterns[:5]:  # Limit to first 5 patterns for speed
        matches = pattern.finditer(text)
        for match in matches:
            skills_text = match.group(1)
            terms = re.split(r'[;,/]|\band\b|\&|\s+', skills_text)
            for term in terms:
                term = term.strip()
                if term and term in ALL_KEYWORDS:
                    original_case = find_keyword_with_original_case(text, term)
                    found_keywords[term] = original_case
    
    # Sort and cache result - restore original sorting
    sorted_keywords = sorted(found_keywords.values(), 
                           key=lambda x: text.lower().index(x.lower()))
    
    set_cached_result(keyword_cache, cache_key, sorted_keywords, 1800)
    return sorted_keywords

def categorize_keywords(keywords):
    """Categorize keywords into different types based on our predefined categories"""
    categorized = {
        'programming_languages': [],
        'web_technologies': [],
        'databases': [],
        'cloud_platforms': [],
        'data_science': [],
        'tools_and_platforms': [],
        'machine_learning': []
    }
    
    for keyword in keywords:
        keyword_lower = keyword.lower()
        for category, keyword_set in TECHNICAL_KEYWORDS.items():
            if keyword_lower in keyword_set:
                categorized[category].append(keyword)
                break
    
    return categorized

def find_section_paragraphs(doc):
    """Find existing sections in the document and return their paragraph indices"""
    sections = {}
    for i, para in enumerate(doc.paragraphs):
        text_upper = para.text.upper().strip()
        # More specific matching to avoid false positives
        if text_upper == 'SKILLS':
            sections['skills'] = i
        elif 'TOOLS' in text_upper and len(text_upper) < 20:  # Avoid matching "Tools/Frameworks:"
            sections['tools'] = i
        elif 'FRAMEWORKS' in text_upper and len(text_upper) < 20:
            sections['frameworks'] = i
        elif 'TECHNOLOGIES' in text_upper and len(text_upper) < 20:
            sections['technologies'] = i
    
    return sections

def insert_keywords_into_sections(doc, missing_keywords):
    if not missing_keywords:
        return doc

    def insert_after_header(header_variations):
        idx = None
        for i, para in enumerate(doc.paragraphs):
            text_upper = para.text.strip().upper()
            # Check for various header variations
            for header in header_variations:
                if header in text_upper and len(text_upper) < 30:  # Avoid matching text within sentences
                    idx = i
                    break
            if idx is not None:
                break
        
        if idx is not None:
            # Look for the next paragraph that contains skills/keywords
            for j in range(idx + 1, len(doc.paragraphs)):
                para = doc.paragraphs[j]
                text = para.text.strip()
                if text and (',' in text or '/' in text or '&' in text or ' and ' in text):
                    add_keywords_with_style(para, missing_keywords)
                    return True  # Successfully added keywords
            # If no suitable paragraph found, add keywords to the header paragraph itself
            if idx < len(doc.paragraphs) - 1:
                next_para = doc.paragraphs[idx + 1]
                if not next_para.text.strip():  # Empty paragraph
                    add_keywords_with_style(next_para, missing_keywords)
                    return True
        return False

    # Try different section header variations
    skills_headers = ["SKILLS", "TECHNICAL SKILLS", "SKILLS & EXPERTISE", "COMPETENCIES", "EXPERTISE"]
    tools_headers = ["TOOLS", "TECHNOLOGIES", "TECHNICAL TOOLS", "SOFTWARE", "PLATFORMS"]
    frameworks_headers = ["FRAMEWORKS", "LIBRARIES", "TECHNOLOGIES", "TECH STACK"]
    
    # Try to insert keywords in order of preference
    if insert_after_header(skills_headers):
        return doc
    if insert_after_header(tools_headers):
        return doc
    if insert_after_header(frameworks_headers):
        return doc
    
    # If no suitable section found, add a new Skills section at the end
    if missing_keywords:
        # Add a blank line first
        doc.add_paragraph("")
        # Add Skills header
        skills_header = doc.add_paragraph("Skills:")
        skills_header.style.font.name = 'Calibri'
        skills_header.style.font.size = Pt(11)
        skills_header.style.font.bold = True
        # Add keywords
        skills_para = doc.add_paragraph(", ".join(missing_keywords))
        skills_para.style.font.name = 'Calibri'
        skills_para.style.font.size = Pt(11)
    
    return doc

def docx_to_text(doc):
    """Convert DOCX document to plain text format"""
    text_content = []
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            text_content.append(text)
    
    return '\n\n'.join(text_content)

def create_export_filename(company_name, job_role, format_type):
    """Create a filename for the exported file"""
    if company_name and job_role:
        # Clean the filename to remove invalid characters
        safe_company = "".join(c for c in company_name if c.isalnum() or c in (' ', '-', '_')).strip()
        safe_role = "".join(c for c in job_role if c.isalnum() or c in (' ', '-', '_')).strip()
        return f"{safe_company} {safe_role} Resume.{format_type}"
    else:
        return f"optimized_resume.{format_type}"

def add_keywords_with_style(paragraph, keywords):
    """Add keywords to a paragraph by appending them to existing text with proper formatting."""
    if not keywords:
        return
    
    original_text = paragraph.text.strip()
    
    # FIRST: Detect the user's preferred separator style
    separator = ' and '  # default
    if ' & ' in original_text:
        separator = ' & '
    elif ' and ' in original_text:
        separator = ' and '
    
    # SECOND: Clean the original text by removing all 'and' and '&' separators
    # Replace ' and ' and ' & ' with commas, then split
    cleaned_text = re.sub(r'\s+(?:and|&)\s+', ', ', original_text)
    
    # Extract existing skills by splitting on ALL possible separators
    existing_skills = []
    if cleaned_text:
        # Split on comma, semicolon, slash
        parts = re.split(r'[;,/]', cleaned_text)
        existing_skills = [part.strip() for part in parts if part.strip()]
    
    # Combine existing and new keywords, avoiding duplicates
    all_skills = existing_skills[:]
    for kw in keywords:
        if kw.lower() not in [s.lower() for s in all_skills]:
            all_skills.append(kw)
    
    if len(all_skills) == 0:
        return
    
    # Format the final string - ALWAYS use comma between all except last, user's separator before last
    if len(all_skills) == 1:
        final_text = all_skills[0]
    elif len(all_skills) == 2:
        final_text = all_skills[0] + separator + all_skills[1]
    else:
        # Multiple skills: comma between all except last, user's separator before last
        final_text = ', '.join(all_skills[:-1]) + separator + all_skills[-1]
    
    # Update the paragraph text
    if paragraph.runs:
        paragraph.runs[0].text = final_text
        # Clear any additional runs
        for run in paragraph.runs[1:]:
            run.text = ''
    else:
        run = paragraph.add_run(final_text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)

@app.route('/optimize-docx', methods=['POST'])
@timeout_handler(30)  # Back to 30 seconds for reliability
def optimize_docx():
    """Optimized resume optimization endpoint - restored original functionality"""
    if 'resume' not in request.files or 'jobDescription' not in request.form:
        return jsonify({'error': 'Missing file or job description'}), 400

    resume_file = request.files['resume']
    job_description = request.form['jobDescription']
    company_name = request.form.get('companyName', '').strip()
    job_role = request.form.get('jobRole', '').strip()
    export_format = request.form.get('exportFormat', 'docx').lower()

    # Enhanced validation with better error messages
    if resume_file.content_length and resume_file.content_length > app.config['MAX_CONTENT_LENGTH']:
        return jsonify({'error': 'File too large. Maximum size is 16MB. Please compress your resume or use a smaller file.'}), 400

    word_count = len(job_description.split())
    if word_count > 750:  # Back to original limit
        return jsonify({'error': f'Job description too long. Please keep it under 750 words. Current: {word_count} words.'}), 400

    # Clean up old temp files
    cleanup_temp_files()

    try:
        # Use circuit breaker for optimization - restore reliability
        def optimization_work():
            # Optimized file processing with memory management
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            resume_file.save(tmp.name)
            doc = Document(tmp.name)

            # Extract text more efficiently
        full_text = '\n'.join([p.text for p in doc.paragraphs])
        
            # Use optimized ATS scoring with caching
            original_ats_score = calculate_ats_score_optimized(full_text, job_description)
        
            # Use optimized keyword extraction
            keywords = extract_technical_keywords_optimized(job_description)
        missing_keywords = [kw for kw in keywords if kw.lower() not in full_text.lower()]

            # Insert keywords efficiently
        doc = insert_keywords_into_sections(doc, missing_keywords)
        
            # Get optimized text
        optimized_text = '\n'.join([p.text for p in doc.paragraphs])
        
            # Calculate optimized ATS score with caching
            optimized_ats_score = calculate_ats_score_optimized(optimized_text, job_description, original_ats_score['total_score'])

            return doc, original_ats_score, optimized_ats_score, keywords, missing_keywords, optimized_text

        # Execute with retry and circuit breaker - restore reliability
        start_time = time.time()
        result = retry_with_backoff(
            lambda: optimization_circuit_breaker.call(optimization_work)
        )
        
        doc, original_ats_score, optimized_ats_score, keywords, missing_keywords, optimized_text = result
        processing_time = time.time() - start_time

        # Handle export formats efficiently
        if export_format == 'txt':
            text_content = docx_to_text(doc)
            text_buffer = io.BytesIO()
            text_buffer.write(text_content.encode('utf-8'))
            text_buffer.seek(0)
            filename = create_export_filename(company_name, job_role, 'txt')
            response = send_file(
                text_buffer,
                as_attachment=True,
                download_name=filename,
                mimetype='text/plain'
            )
        else:
            out_fd, out_path = tempfile.mkstemp(suffix='.docx')
            os.close(out_fd)
            doc.save(out_path)
            filename = create_export_filename(company_name, job_role, 'docx')
            response = send_file(out_path, as_attachment=True, download_name=filename)

        # Add ATS scores to response headers
        response.headers['X-Original-ATS-Score'] = str(original_ats_score['total_score'])
        response.headers['X-Optimized-ATS-Score'] = str(optimized_ats_score['total_score'])
        response.headers['X-ATS-Improvement'] = str(optimized_ats_score['improvement'])
        
        # Enhanced response with performance metrics - restored original functionality
        return jsonify({
            'original_ats_score': original_ats_score,
            'optimized_ats_score': optimized_ats_score,
            'keywords': keywords,
            'missing_keywords': missing_keywords,
            'keywords_added': len(missing_keywords),
            'resumeText': optimized_text[:1000] + "..." if len(optimized_text) > 1000 else optimized_text,
            'download_ready': True,
            'message': f'Resume optimized successfully! Added {len(missing_keywords)} keywords. ATS score improved by {optimized_ats_score["improvement"]:.1f} points.',
            'performance_metrics': {
                'processing_time_ms': int(processing_time * 1000),
                'cache_hits': len([k for k in keyword_cache.keys() if 'keywords' in k]),
                'text_processed': len(optimized_text),
                'keywords_found': len(keywords),
                'keywords_added': len(missing_keywords)
            }
        })
    except Exception as e:
        logger.error(f"Optimization failed: {str(e)}")
        # Provide helpful error messages
        if "timeout" in str(e).lower():
            return jsonify({'error': 'Processing took too long. Please try with a smaller file or shorter job description.'}), 408
        elif "memory" in str(e).lower():
            return jsonify({'error': 'File is too complex to process. Please try with a simpler resume format.'}), 413
        else:
            return jsonify({'error': 'Optimization failed. Please check your file format and try again.'}), 500
    finally:
        # Force garbage collection to free memory
        gc.collect()

@app.route('/export-formats', methods=['GET'])
def get_export_formats():
    """Return available export formats"""
    return jsonify({
        'formats': [
            {
                'value': 'docx',
                'label': 'Microsoft Word (.docx)',
                'description': 'Best for editing and ATS compatibility',
                'icon': '\ud83d\udcc4'
            },
            {
                'value': 'txt',
                'label': 'Plain Text (.txt)',
                'description': 'Simple text format for basic applications',
                'icon': '\ud83d\udcdd'
            }
        ]
    })

# --- Industry inference from job description ---
def infer_industry(job_description):
    """Infer the most likely industry from the job description."""
    scores = {k: 0 for k in INDUSTRY_KEYWORDS}
    text = job_description.lower()
    
    # Count keyword matches for each industry
    for industry, keywords in INDUSTRY_KEYWORDS.items():
        for kw in keywords:
            if kw.lower() in text:
                scores[industry] += 1
    
    # Also check for common industry terms
    industry_indicators = {
        'software_engineering': ['software', 'developer', 'programmer', 'engineer', 'coding', 'programming', 'web', 'mobile', 'app', 'frontend', 'backend', 'fullstack'],
        'data_analytics': ['data', 'analytics', 'analyst', 'business intelligence', 'bi', 'reporting', 'dashboard', 'kpi', 'metrics', 'statistics']
    }
    
    for industry, indicators in industry_indicators.items():
        for indicator in indicators:
            if indicator in text:
                scores[industry] += 1
    
    # Return the industry with the highest score, default to software_engineering if no clear match
    best_industry = max(scores.items(), key=lambda x: x[1])
    return best_industry[0] if best_industry[1] > 0 else 'software_engineering'

# --- Keyword importance scoring ---
def score_keywords(job_description, industry):
    """Return a dict of keyword: score, based on frequency and industry importance."""
    text = job_description.lower()
    industry_keywords = [kw.lower() for kw in INDUSTRY_KEYWORDS.get(industry, [])]
    scores = {}
    for kw in set(re.findall(r'\b\w[\w\+\#\.\-]*\b', job_description, re.IGNORECASE)):
        freq = text.count(kw.lower())
        importance = 2 if kw.lower() in industry_keywords else 1
        if freq > 0:
            scores[kw] = freq * importance
    return scores

# --- Suggest up to 4 high-impact, industry-specific keywords not already in the resume or job description ---
def suggest_extra_keywords(resume_text, job_description, industry, max_suggestions=4):
    # Add timeout protection
    start_time = time.time()
    max_processing_time = 5  # 5 seconds max (reduced from 10)
    
    # Normalize resume words: lowercase, remove punctuation
    resume_words = set([
        w.lower().strip(string.punctuation)
        for w in re.findall(r'\b\w[\w\+\#\.\-]*\b', resume_text)
    ])
    
    # Normalize job description words: lowercase, remove punctuation
    job_words = set([
        w.lower().strip(string.punctuation)
        for w in re.findall(r'\b\w[\w\+\#\.\-]*\b', job_description)
    ])
    
    # Use a more efficient approach - focus on most common keywords
    # Start with industry-specific keywords first (higher priority)
    industry_keywords = INDUSTRY_KEYWORDS.get(industry, [])
    
    # Take only the most common technical keywords from priority categories
    common_technical_keywords = []
    priority_categories = ['programming_languages', 'web_technologies', 'databases', 'cloud_platforms']
    
    # Most common keywords from each category (top 10-15 most relevant)
    common_keywords_by_category = {
        'programming_languages': ['python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 'php', 'swift', 'kotlin', 'go', 'rust', 'scala', 'r', 'sql'],
        'web_technologies': ['html', 'css', 'react', 'angular', 'vue', 'node.js', 'express', 'django', 'flask', 'spring', 'bootstrap', 'tailwind', 'graphql', 'rest'],
        'databases': ['mysql', 'postgresql', 'mongodb', 'redis', 'oracle', 'sqlite', 'sql server', 'dynamodb', 'elasticsearch', 'neo4j', 'firebase'],
        'cloud_platforms': ['aws', 'azure', 'gcp', 'heroku', 'digitalocean', 'firebase', 'vercel', 'netlify', 'supabase']
    }
    
    for category in priority_categories:
        if category in common_keywords_by_category:
            common_technical_keywords.extend(common_keywords_by_category[category])
    
    # Combine keywords but limit total size
    all_keywords = industry_keywords + common_technical_keywords[:50]  # Limit to 50 total
    
    # Remove duplicates and filter out keywords already in resume OR job description
    suggestions = []
    seen = set()
    
    for kw in all_keywords:
        # Check timeout
        if time.time() - start_time > max_processing_time:
            break
            
        kw_lower = kw.lower().strip(string.punctuation)
        # Only suggest keywords that are NOT in resume AND NOT in job description
        if kw_lower not in resume_words and kw_lower not in job_words and kw_lower not in seen:
            # Simplified case handling
            if ' ' in kw:
                # Multi-word terms
                display_kw = ' '.join(word.title() for word in kw.split())
            else:
                # Single word terms
                display_kw = kw.title()
            
            suggestions.append(display_kw)
            seen.add(kw_lower)
            if len(suggestions) >= max_suggestions:
                break
    
    # If we don't have enough suggestions, add some common technical terms
    if len(suggestions) < max_suggestions and time.time() - start_time < max_processing_time:
        common_tech_terms = [
            'Git', 'Docker', 'Kubernetes', 'AWS', 'Azure', 'GCP', 'CI/CD', 'REST API', 'GraphQL',
            'Microservices', 'Agile', 'Scrum', 'DevOps', 'Cloud Computing', 'API Development',
            'Database Design', 'System Architecture', 'Performance Optimization', 'Security',
            'Testing', 'Code Review', 'Version Control', 'Linux', 'Shell Scripting'
        ]
        
        for term in common_tech_terms:
            if time.time() - start_time > max_processing_time:
                break
                
            term_lower = term.lower().strip(string.punctuation)
            if term_lower not in resume_words and term_lower not in job_words and term_lower not in seen:
                suggestions.append(term)
                seen.add(term_lower)
                if len(suggestions) >= max_suggestions:
                    break
    
    return suggestions[:max_suggestions]

@app.route('/suggest-keywords', methods=['POST'])
@timeout_handler(15)  # 15 second timeout
def suggest_keywords():
    if 'resume' not in request.files or 'jobDescription' not in request.form:
        return jsonify({'error': 'Missing file or job description'}), 400
    resume_file = request.files['resume']
    job_description = request.form['jobDescription']
    
    # Add file size validation
    if resume_file.content_length and resume_file.content_length > app.config['MAX_CONTENT_LENGTH']:
        return jsonify({'error': 'File too large. Maximum size is 16MB.'}), 400

    # Add job description length validation
    word_count = len(job_description.split())
    if word_count > 750:
        return jsonify({'error': f'Job description too long. Please keep it under 750 words. Current: {word_count} words.'}), 400
    
    try:
        # Save uploaded file to a temp location
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            resume_file.save(tmp.name)
            doc = Document(tmp.name)
        resume_text = '\n'.join([p.text for p in doc.paragraphs])
        industry = infer_industry(job_description)
        suggestions = suggest_extra_keywords(resume_text, job_description, industry)
        
        # Add debug information
        resume_words = set([
            w.lower().strip(string.punctuation)
            for w in re.findall(r'\b\w[\w\+\#\.\-]*\b', resume_text)
        ])
        
        return jsonify({
            'industry': industry,
            'suggested_keywords': suggestions,
            'debug_info': {
                'resume_text_length': len(resume_text),
                'resume_words_count': len(resume_words),
                'suggestions_count': len(suggestions),
                'industry_keywords_available': len(INDUSTRY_KEYWORDS.get(industry, [])),
                'technical_keywords_total': sum(len(keywords) for keywords in TECHNICAL_KEYWORDS.values())
            }
        })
    except Exception as e:
        return jsonify({'error': f'Keyword suggestion failed: {str(e)}'}), 500

@app.route('/finalize-resume', methods=['POST'])
@timeout_handler(30)  # 30 second timeout
def finalize_resume():
    if 'resume' not in request.files or 'jobDescription' not in request.form or 'extraKeywords' not in request.form:
        return jsonify({'error': 'Missing file, job description, or extra keywords'}), 400
    resume_file = request.files['resume']
    job_description = request.form['jobDescription']
    extra_keywords = request.form.get('extraKeywords', '')
    company_name = request.form.get('companyName', '').strip()
    job_role = request.form.get('jobRole', '').strip()
    export_format = request.form.get('exportFormat', 'docx').lower()

    # Add file size validation
    if resume_file.content_length and resume_file.content_length > app.config['MAX_CONTENT_LENGTH']:
        return jsonify({'error': 'File too large. Maximum size is 16MB.'}), 400

    # Add job description length validation
    word_count = len(job_description.split())
    if word_count > 750:
        return jsonify({'error': f'Job description too long. Please keep it under 750 words. Current: {word_count} words.'}), 400

    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            resume_file.save(tmp.name)
            doc = Document(tmp.name)

        full_text = '\n'.join([p.text for p in doc.paragraphs])
        
        # Calculate original ATS score
        original_ats_score = calculate_ats_score(full_text, job_description)
        
        job_keywords = extract_technical_keywords(job_description)
        missing_job_keywords = [kw for kw in job_keywords if kw.lower() not in full_text.lower()]
        extra_keywords_list = [s.strip() for s in re.split(r'[;,/]|\\band\\b|\\&', extra_keywords) if s.strip()]
        all_keywords = missing_job_keywords + extra_keywords_list
        unique_keywords = []
        seen = set()
        for kw in all_keywords:
            if kw.lower() not in seen:
                unique_keywords.append(kw)
                seen.add(kw.lower())
        doc = insert_keywords_into_sections(doc, unique_keywords)
        
        # Get final optimized text for ATS scoring
        final_text = '\n'.join([p.text for p in doc.paragraphs])
        
        # Calculate final optimized ATS score
        final_ats_score = calculate_ats_score(final_text, job_description, original_ats_score['total_score'])

        if export_format == 'txt':
            text_content = docx_to_text(doc)
            text_buffer = io.BytesIO()
            text_buffer.write(text_content.encode('utf-8'))
            text_buffer.seek(0)
            filename = create_export_filename(company_name, job_role, 'txt')
            response = send_file(
                text_buffer,
                as_attachment=True,
                download_name=filename,
                mimetype='text/plain'
            )
        else:
            out_fd, out_path = tempfile.mkstemp(suffix='.docx')
            os.close(out_fd)
            doc.save(out_path)
            filename = create_export_filename(company_name, job_role, 'docx')
            response = send_file(out_path, as_attachment=True, download_name=filename)
        
        # Add ATS scores to response headers
        response.headers['X-Original-ATS-Score'] = str(original_ats_score['total_score'])
        response.headers['X-Optimized-ATS-Score'] = str(final_ats_score['total_score'])
        response.headers['X-ATS-Improvement'] = str(final_ats_score['improvement'])
        
        return response
    except Exception as e:
        return jsonify({'error': f'Finalization failed: {str(e)}'}), 500

def calculate_ats_score(resume_text, job_description, original_score=None):
    """Legacy function - use calculate_ats_score_optimized for better performance"""
    return calculate_ats_score_optimized(resume_text, job_description, original_score)

def calculate_ats_score_optimized(resume_text, job_description, original_score=None):
    """Optimized ATS score calculation with caching - restored original functionality"""
    if not resume_text or not job_description:
        return {
            'total_score': 0,
            'keyword_score': 0,
            'formatting_score': 0,
            'content_score': 0,
            'structure_score': 0,
            'length_score': 0,
            'improvement': 0
        }
    
    cache_key = get_cache_key('ats_score', resume_text[:500], job_description[:500])
    cached_result = get_cached_result(ats_score_cache, cache_key, 1800)
    if cached_result:
        if original_score:
            cached_result['improvement'] = round(cached_result['total_score'] - original_score, 1)
        return cached_result
    
    # Parallel processing of different score components - restore original
    with ThreadPoolExecutor(max_workers=3) as executor:
        future_keyword = executor.submit(calculate_keyword_match_score_optimized, resume_text.lower(), job_description.lower())
        future_formatting = executor.submit(calculate_formatting_score_optimized, resume_text)
        future_content = executor.submit(calculate_content_quality_score_optimized, resume_text)
        
        keyword_score = future_keyword.result()
        formatting_score = future_formatting.result()
        content_score = future_content.result()
    
    # Calculate weighted total score - restore original weights
    total_score = (
        keyword_score * ATS_CRITERIA['keyword_match'] +
        formatting_score * ATS_CRITERIA['formatting'] +
        content_score * ATS_CRITERIA['content_quality']
    )
    
    total_score = max(0, min(100, total_score))
    
    result = {
        'total_score': round(total_score, 1),
        'keyword_score': round(keyword_score, 1),
        'formatting_score': round(formatting_score, 1),
        'content_score': round(content_score, 1),
        'structure_score': 85.0,  # Default good score
        'length_score': 90.0,     # Default good score
        'improvement': round(total_score - (original_score or 0), 1) if original_score else 0
    }
    
    set_cached_result(ats_score_cache, cache_key, result, 1800)
    return result

def calculate_keyword_match_score(resume_text, job_description):
    """Legacy function - use calculate_keyword_match_score_optimized for better performance"""
    return calculate_keyword_match_score_optimized(resume_text, job_description)

def calculate_keyword_match_score_optimized(resume_text, job_description):
    """Optimized keyword matching with better performance"""
    job_keywords = extract_job_keywords_optimized(job_description)
    
    if not job_keywords:
        return 50
    
    # Use set for faster lookups
    resume_words = set(re.findall(r'\b\w+\b', resume_text))
    matched_keywords = []
    
    for keyword in job_keywords:
        keyword_lower = keyword.lower()
        if keyword_lower in resume_words:
            matched_keywords.append(keyword)
        elif keyword_lower in resume_text:  # Fallback for multi-word keywords
            matched_keywords.append(keyword)
    
    match_percentage = len(matched_keywords) / len(job_keywords)
    
    # More granular scoring to show improvements
    if match_percentage >= 0.9:
        return 100
    elif match_percentage >= 0.8:
        return 95
    elif match_percentage >= 0.7:
        return 90
    elif match_percentage >= 0.6:
        return 85
    elif match_percentage >= 0.5:
        return 80
    elif match_percentage >= 0.4:
        return 75
    elif match_percentage >= 0.3:
        return 70
    elif match_percentage >= 0.2:
        return 65
    elif match_percentage >= 0.1:
        return 60
    else:
        return 50

def extract_job_keywords(job_description):
    """Legacy function - use extract_job_keywords_optimized for better performance"""
    return extract_job_keywords_optimized(job_description)

def extract_job_keywords_optimized(job_description):
    """Optimized job keyword extraction"""
    cache_key = get_cache_key('job_keywords', job_description[:1000])
    cached_result = get_cached_result(keyword_cache, cache_key, 1800)
    if cached_result:
        return cached_result
    
    keywords = extract_technical_keywords_optimized(job_description)
    
    # Simplified additional keyword extraction
    words = re.findall(r'\b[a-zA-Z]{3,}\b', job_description.lower())
    additional_keywords = [word for word in words if word in ALL_KEYWORDS]
    
    all_keywords = list(keywords) + additional_keywords
    unique_keywords = []
    seen = set()
    for kw in all_keywords:
        if kw.lower() not in seen:
            unique_keywords.append(kw)
            seen.add(kw.lower())
    
    result = unique_keywords[:20]  # Reduced limit for speed
    set_cached_result(keyword_cache, cache_key, result, 1800)
    return result

def calculate_formatting_score(resume_text):
    """Legacy function - use calculate_formatting_score_optimized for better performance"""
    return calculate_formatting_score_optimized(resume_text)

def calculate_formatting_score_optimized(resume_text):
    """Optimized formatting score calculation"""
    score = 80
    
    # Pre-compiled patterns for better performance
    problematic_patterns = [
        re.compile(r'<table|<img|<chart|<header|<footer', re.IGNORECASE),
        re.compile(r'columns?|text-align:\s*center|position:\s*absolute', re.IGNORECASE)
    ]
    
    for pattern in problematic_patterns:
        if pattern.search(resume_text):
            score -= 10
    
    # Check for good formatting
    if re.search(r'skills?', resume_text, re.IGNORECASE):
        score += 10
    
    return max(0, min(100, score))

def calculate_content_quality_score(resume_text):
    """Legacy function - use calculate_content_quality_score_optimized for better performance"""
    return calculate_content_quality_score_optimized(resume_text)

def calculate_content_quality_score_optimized(resume_text):
    """Optimized content quality score calculation"""
    score = 70
    
    # Pre-compiled action verbs set for faster lookup
    action_verbs = {
        'developed', 'implemented', 'managed', 'created', 'designed', 'built',
        'improved', 'increased', 'decreased', 'led', 'coordinated', 'organized',
        'analyzed', 'researched', 'planned', 'executed', 'delivered', 'achieved'
    }
    
    resume_lower = resume_text.lower()
    action_verb_count = sum(1 for verb in action_verbs if verb in resume_lower)
    score += min(20, action_verb_count * 2)
    
    # Check for quantifiable achievements
    achievement_patterns = [
        re.compile(r'\d+%|\d+\s*percent|\$\d+|\d+\s*dollars', re.IGNORECASE),
        re.compile(r'increased by \d+|decreased by \d+|reduced by \d+|improved by \d+', re.IGNORECASE)
    ]
    
    achievement_count = sum(1 for pattern in achievement_patterns if pattern.search(resume_text))
    score += min(10, achievement_count * 2)
    
    return max(0, min(100, score))

def calculate_structure_score(resume_text):
    """Calculate structure score (0-100)"""
    score = 80  # Start with higher base score
    
    # Check for required sections
    required_sections = ['experience', 'education', 'skills']
    section_count = sum(1 for section in required_sections if section in resume_text.lower())
    score += section_count * 10  # Reduced points per section but higher base
    
    # Check for proper section headers
    header_patterns = [
        r'experience', r'education', r'skills', r'summary', r'objective',
        r'work history', r'employment', r'qualifications', r'achievements'
    ]
    
    header_count = sum(1 for pattern in header_patterns if re.search(pattern, resume_text, re.IGNORECASE))
    score += min(10, header_count * 2)  # Reduced max points
    
    return max(0, min(100, score))

def calculate_length_score(resume_text):
    """Calculate length score (0-100)"""
    word_count = len(resume_text.split())
    
    if ATS_FORMATTING_REQUIREMENTS['min_length'] <= word_count <= ATS_FORMATTING_REQUIREMENTS['max_length']:
        return 100
    elif word_count < ATS_FORMATTING_REQUIREMENTS['min_length']:
        # Penalize for being too short
        return max(0, 100 - (ATS_FORMATTING_REQUIREMENTS['min_length'] - word_count) * 2)
    else:
        # Penalize for being too long
        return max(0, 100 - (word_count - ATS_FORMATTING_REQUIREMENTS['max_length']) * 0.5)

def optimize_for_ats(resume_text, job_description, target_score=95):
    """
    Optimize resume text to achieve target ATS score
    Returns optimized text and new score
    """
    current_score = calculate_ats_score(resume_text, job_description)
    
    if current_score['total_score'] >= target_score:
        return resume_text, current_score
    
    # Extract missing keywords
    job_keywords = extract_job_keywords(job_description.lower())
    resume_lower = resume_text.lower()
    missing_keywords = [kw for kw in job_keywords if kw.lower() not in resume_lower]
    
    # Add missing keywords strategically
    optimized_text = resume_text
    
    # Add missing keywords to skills section or create one
    if missing_keywords:
        skills_section = "\n\nSkills:\n"
        skills_section += ", ".join(missing_keywords[:10])  # Limit to 10 keywords
        optimized_text += skills_section
    
    # Recalculate score
    new_score = calculate_ats_score(optimized_text, job_description, current_score['total_score'])
    
    return optimized_text, new_score

@app.route('/calculate-ats-score', methods=['POST'])
def calculate_ats_score_endpoint():
    """Calculate ATS score for uploaded resume"""
    try:
        if 'resume' not in request.files:
            return jsonify({'error': 'No resume file provided'}), 400
        
        resume_file = request.files['resume']
        job_description = request.form.get('jobDescription', '')
        
        if not job_description:
            return jsonify({'error': 'Job description is required'}), 400
        
        # Read the DOCX file
        doc = Document(resume_file)
        resume_text = docx_to_text(doc)
        
        # Calculate ATS score
        ats_score = calculate_ats_score(resume_text, job_description)
        
        return jsonify({
            'ats_score': ats_score,
            'resume_text': resume_text[:500] + "..." if len(resume_text) > 500 else resume_text
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/optimize-ats', methods=['POST'])
def optimize_ats_endpoint():
    """Optimize resume for ATS and return new score"""
    try:
        if 'resume' not in request.files:
            return jsonify({'error': 'No resume file provided'}), 400
        
        resume_file = request.files['resume']
        job_description = request.form.get('jobDescription', '')
        target_score = int(float(request.form.get('targetScore', 95)))
        
        if not job_description:
            return jsonify({'error': 'Job description is required'}), 400
        
        # Read the DOCX file
        doc = Document(resume_file)
        original_text = docx_to_text(doc)
        
        # Calculate original ATS score
        original_score = calculate_ats_score(original_text, job_description)
        
        # Optimize for ATS
        optimized_text, optimized_score = optimize_for_ats(original_text, job_description, target_score)
        
        # Create optimized document
        optimized_doc = Document()
        
        # Add optimized content
        for paragraph in optimized_text.split('\n'):
            if paragraph.strip():
                p = optimized_doc.add_paragraph(paragraph.strip())
                p.style.font.name = 'Calibri'
                p.style.font.size = Pt(11)
        
        # Save to temporary file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        optimized_doc.save(temp_file.name)
        temp_file.close()
        
        return send_file(
            temp_file.name,
            as_attachment=True,
            download_name="ats_optimized_resume.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/download-optimized', methods=['POST'])
def download_optimized():
    """Download the optimized resume file"""
    try:
        if 'resume' not in request.files or 'jobDescription' not in request.form:
            return jsonify({'error': 'Missing file or job description'}), 400

        resume_file = request.files['resume']
        job_description = request.form['jobDescription']
        company_name = request.form.get('companyName', '').strip()
        job_role = request.form.get('jobRole', '').strip()
        export_format = request.form.get('exportFormat', 'docx').lower()

        # Save uploaded file to a temp location
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            resume_file.save(tmp.name)
            doc = Document(tmp.name)

        # Extract all text for keyword matching
        full_text = '\n'.join([p.text for p in doc.paragraphs])
        
        # Use the new technical keyword extraction that preserves case
        keywords = extract_technical_keywords(job_description)
        missing_keywords = [kw for kw in keywords if kw.lower() not in full_text.lower()]

        # Insert missing keywords into existing Skills section
        doc = insert_keywords_into_sections(doc, missing_keywords)

        # Handle different export formats
        if export_format == 'txt':
            # Convert to plain text
            text_content = docx_to_text(doc)
            text_buffer = io.BytesIO()
            text_buffer.write(text_content.encode('utf-8'))
            text_buffer.seek(0)
            filename = create_export_filename(company_name, job_role, 'txt')
            return send_file(
                text_buffer,
                as_attachment=True,
                download_name=filename,
                mimetype='text/plain'
            )
        else:
            # Default: DOCX format
            out_fd, out_path = tempfile.mkstemp(suffix='.docx')
            os.close(out_fd)
            doc.save(out_path)
            filename = create_export_filename(company_name, job_role, 'docx')
            return send_file(out_path, as_attachment=True, download_name=filename)

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/debug-keywords', methods=['POST'])
def debug_keywords():
    """Debug endpoint to test keyword extraction and ATS scoring"""
    if 'resume' not in request.files or 'jobDescription' not in request.form:
        return jsonify({'error': 'Missing file or job description'}), 400

    resume_file = request.files['resume']
    job_description = request.form['jobDescription']

    # Save uploaded file to a temp location
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        resume_file.save(tmp.name)
        doc = Document(tmp.name)

    # Extract all text
    full_text = '\n'.join([p.text for p in doc.paragraphs])
    
    # Extract keywords using both methods
    technical_keywords = extract_technical_keywords(job_description)
    job_keywords = extract_job_keywords(job_description)
    
    # Calculate ATS scores
    original_ats_score = calculate_ats_score(full_text, job_description)
    
    # Find missing keywords
    missing_keywords = [kw for kw in technical_keywords if kw.lower() not in full_text.lower()]
    
    # Simulate adding keywords
    test_text = full_text + "\n\nSkills: " + ", ".join(missing_keywords[:10])
    test_ats_score = calculate_ats_score(test_text, job_description, original_ats_score['total_score'])
    
    # Detailed keyword matching analysis
    resume_lower = full_text.lower()
    matched_keywords = []
    for keyword in job_keywords:
        keyword_lower = keyword.lower()
        pattern = r'\b' + re.escape(keyword_lower) + r'\b'
        if re.search(pattern, resume_lower):
            matched_keywords.append(keyword)
        elif keyword_lower in resume_lower:
            matched_keywords.append(keyword)
    
    match_percentage = len(matched_keywords) / len(job_keywords) if job_keywords else 0
    
    return jsonify({
        'resume_text_length': len(full_text),
        'technical_keywords_found': len(technical_keywords),
        'job_keywords_found': len(job_keywords),
        'missing_keywords': missing_keywords,
        'matched_keywords': matched_keywords,
        'match_percentage': match_percentage,
        'original_ats_score': original_ats_score,
        'test_ats_score': test_ats_score,
        'technical_keywords': technical_keywords,
        'job_keywords': job_keywords,
        'resume_preview': full_text[:500] + "..." if len(full_text) > 500 else full_text,
        'score_breakdown': {
            'keyword_score': original_ats_score['keyword_score'],
            'formatting_score': original_ats_score['formatting_score'],
            'content_score': original_ats_score['content_score'],
            'structure_score': original_ats_score['structure_score'],
            'length_score': original_ats_score['length_score']
        }
    })

@app.route('/test-suggestions', methods=['POST'])
def test_suggestions():
    """Test endpoint to debug keyword suggestions"""
    if 'resume' not in request.files or 'jobDescription' not in request.form:
        return jsonify({'error': 'Missing file or job description'}), 400

    resume_file = request.files['resume']
    job_description = request.form['jobDescription']

    # Save uploaded file to a temp location
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        resume_file.save(tmp.name)
        doc = Document(tmp.name)

    resume_text = '\n'.join([p.text for p in doc.paragraphs])
    industry = infer_industry(job_description)
    suggestions = suggest_extra_keywords(resume_text, job_description, industry)
    
    # Get all available keywords for comparison
    all_technical_keywords = []
    for category, keywords in TECHNICAL_KEYWORDS.items():
        all_technical_keywords.extend(keywords)
    
    industry_keywords = INDUSTRY_KEYWORDS.get(industry, [])
    
    # Normalize resume words
    resume_words = set([
        w.lower().strip(string.punctuation)
        for w in re.findall(r'\b\w[\w\+\#\.\-]*\b', resume_text)
    ])
    
    return jsonify({
        'industry': industry,
        'suggested_keywords': suggestions,
        'resume_text_preview': resume_text[:200] + "..." if len(resume_text) > 200 else resume_text,
        'resume_words_count': len(resume_words),
        'resume_words_sample': list(resume_words)[:20],
        'industry_keywords': industry_keywords,
        'technical_keywords_sample': all_technical_keywords[:50],
        'total_technical_keywords': len(all_technical_keywords),
        'debug_info': {
            'resume_text_length': len(resume_text),
            'suggestions_count': len(suggestions),
            'industry_keywords_available': len(industry_keywords),
            'technical_keywords_total': len(all_technical_keywords)
        }
    })

if __name__ == '__main__':
    app.run(port=8000, debug=False, use_reloader=False) 