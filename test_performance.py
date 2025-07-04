#!/usr/bin/env python3
"""
Performance test script for the resume optimization tool
"""

import requests
import time
import json
from pathlib import Path

# Configuration
API_BASE_URL = "http://localhost:5000"
TEST_RESUME_PATH = "test_resume.docx"  # Create a simple test resume
TEST_JOB_DESCRIPTION = """
Software Engineer position requiring Python, JavaScript, React, Node.js, SQL, AWS, Docker, Git, Agile methodologies, and experience with web development frameworks. 
The ideal candidate should have experience with REST APIs, microservices architecture, and cloud platforms.
"""

def create_test_resume():
    """Create a simple test resume if it doesn't exist"""
    if Path(TEST_RESUME_PATH).exists():
        return
    
    from docx import Document
    from docx.shared import Inches
    
    doc = Document()
    
    # Add content
    doc.add_heading('Test Resume', 0)
    
    doc.add_heading('Experience', level=1)
    p = doc.add_paragraph('Software Engineer at Test Company')
    p.add_run(' (2020-2023)')
    doc.add_paragraph('• Developed web applications using Python and JavaScript')
    doc.add_paragraph('• Worked with React and Node.js frameworks')
    doc.add_paragraph('• Managed databases using SQL and MongoDB')
    
    doc.add_heading('Skills', level=1)
    doc.add_paragraph('Python, JavaScript, React, Node.js, SQL, Git')
    
    doc.add_heading('Education', level=1)
    doc.add_paragraph('Bachelor of Science in Computer Science')
    
    doc.save(TEST_RESUME_PATH)
    print(f"Created test resume: {TEST_RESUME_PATH}")

def test_health():
    """Test health endpoint"""
    try:
        response = requests.get(f"{API_BASE_URL}/health", timeout=5)
        print(f"Health check: {response.status_code} - {response.json()}")
        return response.status_code == 200
    except Exception as e:
        print(f"Health check failed: {e}")
        return False

def test_status():
    """Test status endpoint"""
    try:
        response = requests.get(f"{API_BASE_URL}/status", timeout=5)
        data = response.json()
        print(f"Status: Fast mode = {data.get('fast_mode')}, Active requests = {data.get('active_requests')}")
        return data
    except Exception as e:
        print(f"Status check failed: {e}")
        return None

def test_optimization():
    """Test resume optimization performance"""
    if not Path(TEST_RESUME_PATH).exists():
        print("Test resume not found. Please create a test resume first.")
        return
    
    try:
        with open(TEST_RESUME_PATH, 'rb') as f:
            files = {'resume': f}
            data = {
                'jobDescription': TEST_JOB_DESCRIPTION,
                'companyName': 'Test Company',
                'jobRole': 'Software Engineer',
                'exportFormat': 'docx'
            }
            
            print("Testing resume optimization...")
            start_time = time.time()
            
            response = requests.post(
                f"{API_BASE_URL}/optimize-docx",
                files=files,
                data=data,
                timeout=30
            )
            
            end_time = time.time()
            processing_time = end_time - start_time
            
            if response.status_code == 200:
                result = response.json()
                print(f"✅ Optimization successful!")
                print(f"⏱️  Processing time: {processing_time:.2f} seconds")
                print(f"📊 Original ATS score: {result.get('original_ats_score', {}).get('total_score', 'N/A')}")
                print(f"📈 Optimized ATS score: {result.get('optimized_ats_score', {}).get('total_score', 'N/A')}")
                print(f"🔑 Keywords added: {result.get('keywords_added', 0)}")
                
                if 'performance_metrics' in result:
                    metrics = result['performance_metrics']
                    print(f"⚡ Fast mode: {metrics.get('fast_mode', 'N/A')}")
                    print(f"💾 Cache hits: {metrics.get('cache_hits', 'N/A')}")
                
                return processing_time
            else:
                print(f"❌ Optimization failed: {response.status_code}")
                print(f"Error: {response.text}")
                return None
                
    except Exception as e:
        print(f"❌ Test failed: {e}")
        return None

def test_multiple_requests():
    """Test multiple concurrent requests"""
    print("\nTesting multiple requests...")
    
    import concurrent.futures
    
    def single_request():
        try:
            with open(TEST_RESUME_PATH, 'rb') as f:
                files = {'resume': f}
                data = {
                    'jobDescription': TEST_JOB_DESCRIPTION,
                    'companyName': 'Test Company',
                    'jobRole': 'Software Engineer',
                    'exportFormat': 'docx'
                }
                
                start_time = time.time()
                response = requests.post(
                    f"{API_BASE_URL}/optimize-docx",
                    files=files,
                    data=data,
                    timeout=30
                )
                end_time = time.time()
                
                return {
                    'success': response.status_code == 200,
                    'time': end_time - start_time,
                    'status': response.status_code
                }
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    # Test 3 concurrent requests
    with concurrent.futures.ThreadPoolExecutor(max_workers=3) as executor:
        futures = [executor.submit(single_request) for _ in range(3)]
        results = [future.result() for future in concurrent.futures.as_completed(futures)]
    
    successful = [r for r in results if r['success']]
    print(f"✅ Successful requests: {len(successful)}/3")
    
    if successful:
        avg_time = sum(r['time'] for r in successful) / len(successful)
        print(f"⏱️  Average processing time: {avg_time:.2f} seconds")
        print(f"🚀 Fastest request: {min(r['time'] for r in successful):.2f} seconds")
        print(f"🐌 Slowest request: {max(r['time'] for r in successful):.2f} seconds")

def main():
    """Run all performance tests"""
    print("🚀 Resume Optimization Performance Test")
    print("=" * 50)
    
    # Create test resume
    create_test_resume()
    
    # Test health
    print("\n1. Testing health endpoint...")
    if not test_health():
        print("❌ Health check failed. Make sure the server is running.")
        return
    
    # Test status
    print("\n2. Testing status endpoint...")
    status = test_status()
    if status:
        print(f"✅ Server is operational (uptime: {status.get('uptime', 0):.1f}s)")
    
    # Test single optimization
    print("\n3. Testing single optimization...")
    processing_time = test_optimization()
    
    if processing_time:
        if processing_time < 5:
            print("🎉 Excellent performance! (< 5 seconds)")
        elif processing_time < 10:
            print("👍 Good performance! (< 10 seconds)")
        else:
            print("⚠️  Performance could be improved (> 10 seconds)")
    
    # Test multiple requests
    print("\n4. Testing concurrent requests...")
    test_multiple_requests()
    
    print("\n" + "=" * 50)
    print("🏁 Performance test completed!")

if __name__ == "__main__":
    main() 